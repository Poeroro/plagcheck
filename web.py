"""DoubleCheck FastAPI web server — production version."""
from __future__ import annotations

import json
import secrets
import shutil
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, Response, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.base import BaseHTTPMiddleware

from core import Corpus, PlagEngine
from core.engine import ONNX_HYBRID_PATH

ROOT = Path(__file__).parent
UPLOAD_DIR = ROOT / "uploads"
REPORT_DIR = ROOT / "reports"
UPLOAD_DIR.mkdir(exist_ok=True)
REPORT_DIR.mkdir(exist_ok=True)

APP_VERSION = "0.7.1"
MAX_UPLOAD_MB = 10

# Per-session privacy: each browser gets its own cookie-stamped session ID
# Reports and uploads are namespaced by this prefix so users can only see
# their own documents. No login required — first visit creates a session.
SESSION_COOKIE = "pc_sid"
SESSION_MAX_AGE = 30 * 24 * 3600  # 30 days
SESSION_PREFIX_LEN = 8  # first 8 alnum chars used in filename prefix (48 bits)

app = FastAPI(title="DoubleCheck", version=APP_VERSION)
templates = Jinja2Templates(directory=str(ROOT / "templates"))
app.mount("/static", StaticFiles(directory=str(ROOT / "static")), name="static")


# -----------------------------------------------------------------------
# Session middleware: per-browser cookie-based privacy
# -----------------------------------------------------------------------
class SessionMiddleware(BaseHTTPMiddleware):
    """Attach a session ID to every request via HttpOnly Secure cookie.

    On first visit, a fresh 32-byte URL-safe token is generated and set
    as the pc_sid cookie. The token is also stored on request.state.sid
    so endpoints can read it without re-parsing cookies.
    """

    async def dispatch(self, request: Request, call_next):
        sid = request.cookies.get(SESSION_COOKIE)
        if not _is_valid_sid(sid):
            sid = secrets.token_urlsafe(32)
        request.state.sid = sid
        response = await call_next(request)
        # Always set/refresh the cookie so the session stays alive
        # (response.set_cookie is a no-op if the value matches; the
        # browser updates the expiry either way).
        response.set_cookie(
            key=SESSION_COOKIE,
            value=sid,
            max_age=SESSION_MAX_AGE,
            httponly=True,
            samesite="lax",
            secure=True,
            path="/",
        )
        return response


# Add session middleware (must come after route definitions to ensure
# it wraps everything, but Starlette wraps in reverse so add before mount)
app.add_middleware(SessionMiddleware)

_engine: PlagEngine | None = None
_corpus: Corpus | None = None
_semantic_loaded: bool = False
_ensemble_loaded: bool = False
_cross_encoder_loaded: bool = False
_ai_detector_loaded: bool = False


# -----------------------------------------------------------------------
# Startup: pre-warm the ONNX encoder in a background thread so the first
# user request doesn't pay the 5-10s model-load latency. The encoder is
# the most commonly used one (Standar preset uses semantic), so warming
# it removes the cold-start penalty for the typical user flow.
# -----------------------------------------------------------------------
import asyncio  # noqa: E402
from contextlib import asynccontextmanager  # noqa: E402


def _prewarm_semantic() -> None:
    """Run in a thread; logs progress. Errors are non-fatal."""
    try:
        print("[startup] pre-warming ONNX semantic encoder (background)…", flush=True)
        ensure_semantic_loaded()
        print("[startup] ONNX encoder ready ✓", flush=True)
    except Exception as e:  # noqa: BLE001
        print(f"[startup] semantic pre-warm failed (non-fatal): {e}", flush=True)


def _prewarm_optional() -> None:
    """Best-effort pre-warm of cross-encoder and AI detector. Failures
    are non-fatal — first request will still work, just cold-loads."""
    try:
        print("[startup] pre-warming cross-encoder (background)…", flush=True)
        ensure_cross_encoder_loaded()
        print("[startup] cross-encoder ready ✓", flush=True)
    except Exception as e:  # noqa: BLE001
        print(f"[startup] cross-encoder pre-warm failed (non-fatal): {e}", flush=True)
    try:
        print("[startup] pre-warming AI detector (background)…", flush=True)
        ensure_ai_detector_loaded()
        print("[startup] AI detector ready ✓", flush=True)
    except Exception as e:  # noqa: BLE001
        print(f"[startup] AI detector pre-warm failed (non-fatal): {e}", flush=True)


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup — run all pre-warms in a single background thread so they
    # don't compete with the main event loop. ONNX first (most common),
    # then CE + AI in sequence.
    loop = asyncio.get_running_loop()

    def _all_prewarm() -> None:
        _prewarm_semantic()
        _prewarm_optional()

    loop.run_in_executor(None, _all_prewarm)
    yield
    # Shutdown (nothing to clean up explicitly)


# Re-bind app to use the lifespan (FastAPI doesn't let us set lifespan at
# construction time after add_middleware in some versions — re-create cleanly).
app.router.lifespan_context = lifespan


def get_engine() -> PlagEngine:
    global _engine, _corpus
    if _engine is None:
        _corpus = Corpus(ROOT / "corpus")
        _engine = PlagEngine(_corpus)
    return _engine


def ensure_semantic_loaded() -> None:
    global _semantic_loaded
    eng = get_engine()
    if not _semantic_loaded:
        eng.enable_semantic()
        _semantic_loaded = True


def ensure_ensemble_loaded() -> None:
    global _ensemble_loaded
    eng = get_engine()
    if not _ensemble_loaded:
        eng.enable_ensemble()
        _ensemble_loaded = True


def ensure_cross_encoder_loaded() -> None:
    global _cross_encoder_loaded
    eng = get_engine()
    if not _cross_encoder_loaded:
        eng.enable_cross_encoder()
        _cross_encoder_loaded = True


def ensure_ai_detector_loaded() -> None:
    global _ai_detector_loaded
    if not _ai_detector_loaded:
        from core.ai_detector import _get_detector
        _get_detector()
        _ai_detector_loaded = True


# -----------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------
def risk_label(score: float) -> str:
    if score > 0.30:
        return "TINGGI"
    if score > 0.10:
        return "SEDANG"
    return "RENDAH"


def risk_class(score: float) -> str:
    if score > 0.30:
        return "high"
    if score > 0.10:
        return "med"
    return "low"


def pct_class(score: float) -> str:
    if score > 0.70:
        return "high"
    if score > 0.40:
        return "med"
    return "low"


def ai_verdict_class(verdict: str) -> str:
    return {
        "AI-LIKELY": "AI-LIKELY",
        "MIXED": "MIXED",
        "HUMAN": "HUMAN",
    }.get(verdict, "MIXED")


def humanize_ago(iso: str) -> str:
    try:
        ts = datetime.fromisoformat(iso)
        delta = datetime.utcnow() - ts
        secs = int(delta.total_seconds())
        if secs < 60: return f"{secs} detik lalu"
        if secs < 3600: return f"{secs // 60} menit lalu"
        if secs < 86400: return f"{secs // 3600} jam lalu"
        return f"{secs // 86400} hari lalu"
    except Exception:
        return iso


def humanize_time(iso: str) -> str:
    try:
        ts = datetime.fromisoformat(iso)
        return ts.strftime("%d %b %Y · %H:%M")
    except Exception:
        return iso


def build_id_from_path(path: Path) -> str:
    return path.stem  # <timestamp>_<session_prefix>_<safe_filename>


def _is_valid_sid(sid: str) -> bool:
    """Validate a session ID: at least 16 chars, alnum + - _ only."""
    if not isinstance(sid, str) or len(sid) < 16 or len(sid) > 128:
        return False
    return all(c.isalnum() or c in "-_" for c in sid)


def session_prefix(sid: str) -> str:
    """Filesystem-safe prefix derived from session ID for filename scoping."""
    return "".join(c for c in sid[:SESSION_PREFIX_LEN] if c.isalnum())[:SESSION_PREFIX_LEN]


def get_session_id(request: Request) -> str:
    """Read the session ID set by SessionMiddleware on request.state.

    Endpoints that need the session ID should declare this as a
    dependency. If the middleware did not run (e.g. on a synthetic
    test request), a fresh ID is minted on the fly — but in normal
    production flow it is always populated.
    """
    sid = getattr(request.state, "sid", None)
    if not _is_valid_sid(sid):
        sid = secrets.token_urlsafe(32)
        request.state.sid = sid
    return sid


def list_recent_reports(limit: int = 20, session_id: str = "") -> list[dict]:
    """List reports belonging to the given session only.

    Filenames are scoped as <timestamp>_<session_prefix>_<safe_name>.json,
    so we glob for files starting with the session prefix to enforce
    cross-session isolation at the filesystem level.
    """
    reports = []
    prefix = session_prefix(session_id) if session_id else ""
    if not prefix:
        # No valid session — return nothing rather than leak everything
        return reports
    # Filename format: <timestamp>_<session_prefix>_<name>.json
    # The session prefix is the 2nd underscore-delimited token, so we
    # glob for *_<prefix>_*.json and verify on the loaded result.
    pattern = f"*_{prefix}_*.json"
    for path in sorted(REPORT_DIR.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True):
        # Defense in depth: confirm the second token matches
        stem = path.stem
        parts = stem.split("_", 2)
        if len(parts) < 3 or parts[1] != prefix:
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        report_id = build_id_from_path(path)
        score = data.get("overall_score", 0) * 100
        ai = data.get("ai_detection")
        ai_label = ""
        ai_class = "ai-low"
        if ai and not ai.get("error"):
            ai_pct = ai.get("ai_probability", 0) * 100
            ai_label = f"{ai_pct:.0f}%"
            ai_class = "ai-high" if ai_pct > 60 else "ai-med" if ai_pct > 30 else "ai-low"
        raw_filename = data.get("document_name", "") or "?"
        # Skip orphan/empty entries
        if not raw_filename or raw_filename == "?" or data.get("total_paragraphs", 0) == 0:
            continue
        # Clean up filename: strip "<timestamp>_<session_prefix>_" prefix
        display_filename = raw_filename
        if "_" in display_filename:
            parts = raw_filename.split("_", 2)
            # New format: <ts>_<sid8>_<name> — strip first two
            if len(parts) >= 3 and parts[0].isdigit() and len(parts[0]) >= 10:
                display_filename = parts[2]
            # Old format: <ts>_<name> — strip first
            elif len(parts) >= 2 and parts[0].isdigit() and len(parts[0]) >= 10:
                display_filename = parts[1]
        filetype = (raw_filename.split(".")[-1] or "txt").lower()
        if filetype not in ("pdf", "docx", "txt"):
            filetype = "txt"
        reports.append({
            "id": report_id,
            "filename": display_filename,
            "raw_filename": raw_filename,
            "score": f"{score:.1f}",
            "score_class": risk_class(data.get("overall_score", 0)),
            "elapsed": f"{data.get('elapsed_seconds', 0):.1f}",
            "paragraphs": data.get("total_paragraphs", 0),
            "ai_label": ai_label,
            "ai_class": ai_class,
            "filetype": filetype,
            "finished_at_human": humanize_time(data.get("finished_at", "")),
            "ago": humanize_ago(data.get("finished_at", "")),
        })
    return reports[:limit]


def load_report(report_id: str, session_id: str = "") -> tuple[dict, Path]:
    """Return parsed report data + path. Raises 404 if not found or
    not owned by the given session (privacy enforcement)."""
    # Defense in depth: reject obvious cross-session access
    prefix = session_prefix(session_id) if session_id else ""
    if not prefix:
        raise HTTPException(404, "Report not found")
    # Filename format: <timestamp>_<session_prefix>_<name>
    # The session prefix is the 2nd underscore-delimited token.
    parts = report_id.split("_", 2)
    if len(parts) < 3 or parts[1] != prefix:
        raise HTTPException(404, "Report not found")
    matches = list(REPORT_DIR.glob(f"{report_id}.*"))
    if not matches:
        raise HTTPException(404, "Report not found")
    json_path = next((m for m in matches if m.suffix == ".json"), None)
    if not json_path:
        raise HTTPException(404, "Report JSON not found")
    # Final filesystem-level check: filename must contain the session prefix
    if prefix not in json_path.name:
        raise HTTPException(404, "Report not found")
    # And the second token must match exactly
    json_parts = json_path.stem.split("_", 2)
    if len(json_parts) < 3 or json_parts[1] != prefix:
        raise HTTPException(404, "Report not found")
    try:
        data = json.loads(json_path.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(500, f"Failed to read report: {e}")
    return data, json_path


def enrich_report(data: dict) -> dict:
    """Add derived UI properties to report data."""
    score = data.get("overall_score", 0)
    data["risk_label"] = risk_label(score)
    data["risk_class"] = risk_class(score)
    for m in data.get("matches", []):
        m["pct_class"] = pct_class(m.get("score", 0))
    ai = data.get("ai_detection")
    if ai and not ai.get("error"):
        ai["verdict_class"] = ai_verdict_class(ai.get("verdict", "MIXED"))
    # Clean up filename for display
    # Filename format on disk: <timestamp>_<session_prefix>_<original_name>
    # Example: "1781338150646_hotUU3aC_Aril Mira Rahayu - Revisi Tugas 3.json"
    raw = data.get("document_name", "")
    parts = raw.split("_", 2)
    if len(parts) == 3 and parts[0].isdigit() and len(parts[0]) >= 10 and parts[1]:
        clean_name = parts[2]
    elif "_" in raw and raw.split("_", 1)[0].isdigit():
        clean_name = raw.split("_", 1)[1]
    else:
        clean_name = raw
    # Strip common document extensions from the display name for cleanness
    for ext in (".txt", ".docx", ".doc", ".pdf", ".md", ".rtf"):
        if clean_name.lower().endswith(ext):
            clean_name = clean_name[: -len(ext)]
            break
    data["display_name"] = clean_name or raw
    # Compute citation reduction if not present
    cs = data.get("citation_stats") or {}
    if cs and "reduction_pct" not in cs:
        orig = cs.get("original_chars", 0)
        clean = cs.get("cleaned_chars", 0)
        if orig > 0:
            cs["reduction_pct"] = (1 - clean / orig) * 100
        else:
            cs["reduction_pct"] = 0
        data["citation_stats"] = cs
    # Human-readable timestamps
    from datetime import datetime as _dt
    for field in ("started_at", "finished_at"):
        val = data.get(field, "")
        if val and "T" in str(val):
            try:
                dt = _dt.fromisoformat(str(val).replace("Z", ""))
                bulan = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
                         "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
                data[field + "_human"] = f"{dt.day} {bulan[dt.month]} {dt.year}, {dt.strftime('%H:%M')}"
            except Exception:
                data[field + "_human"] = val
    return data


# -----------------------------------------------------------------------
# Routes
# -----------------------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def landing(
    request: Request,
    sid: str = Depends(get_session_id),
) -> HTMLResponse:
    eng = get_engine()
    recent = list_recent_reports(limit=5, session_id=sid)
    return templates.TemplateResponse(
        request,
        "landing.html",
        {
            "version": APP_VERSION,
            "corpus_size": f"{len(eng.corpus):,}",
            "recent": recent,
        },
    )


@app.get("/riwayat", response_class=HTMLResponse)
async def riwayat(
    request: Request,
    sid: str = Depends(get_session_id),
) -> HTMLResponse:
    eng = get_engine()
    recent = list_recent_reports(limit=50, session_id=sid)
    return templates.TemplateResponse(
        request,
        "riwayat.html",
        {
            "version": APP_VERSION,
            "corpus_size": f"{len(eng.corpus):,}",
            "recent": recent,
            "total_reports": len(recent),
        },
    )


@app.get("/r/{report_id}", response_class=HTMLResponse)
async def report_page(
    request: Request,
    report_id: str,
    sid: str = Depends(get_session_id),
) -> HTMLResponse:
    data, _ = load_report(report_id, session_id=sid)
    data = enrich_report(data)
    total = len(list(REPORT_DIR.glob(f"*_{session_prefix(sid)}_*.json")))
    return templates.TemplateResponse(
        request,
        "results.html",
        {
            "version": APP_VERSION,
            "r": data,
            "total_reports": total,
        },
    )


@app.get("/r/{report_id}/json")
async def report_json(
    report_id: str,
    sid: str = Depends(get_session_id),
) -> JSONResponse:
    data, _ = load_report(report_id, session_id=sid)
    return JSONResponse(data)


@app.get("/r/{report_id}/html")
async def report_html(
    request: Request,
    report_id: str,
    sid: str = Depends(get_session_id),
) -> HTMLResponse:
    data, json_path = load_report(report_id, session_id=sid)
    data = enrich_report(data)
    eng = get_engine()
    # Reconstruct a CheckResult for the legacy report builder
    from core.report import CheckResult, CitationStats, Match
    result = CheckResult(
        document_name=data.get("document_name", ""),
        total_paragraphs=data.get("total_paragraphs", 0),
        matches=[Match(**{k: m.get(k) for k in
            ["query_text","matched_text","score","source_title","source_url","source_id","source_type","preview_html"]})
            for m in data.get("matches", [])],
        overall_score=data.get("overall_score", 0),
        flagged_paragraphs=data.get("flagged_paragraphs", 0),
        started_at=data.get("started_at", ""),
        finished_at=data.get("finished_at", ""),
        elapsed_seconds=data.get("elapsed_seconds", 0),
        corpus_size=data.get("corpus_size", 0),
        citation_stats=CitationStats(**{k: v for k, v in (data.get("citation_stats") or {}).items()
            if k in CitationStats.__dataclass_fields__}),
    )
    html = eng.report_html(result)
    return HTMLResponse(
        html,
        headers={"Content-Disposition": f'attachment; filename="{report_id}.html"'},
    )


# ---------- Plain-text export ----------

def _report_text(data: dict) -> str:
    """Build a clean plain-text report from enriched data."""
    name = data.get("display_name") or data.get("document_name", "dokumen")
    score = data.get("overall_score", 0) * 100
    flagged = data.get("flagged_paragraphs", 0)
    total = data.get("total_paragraphs", 0)
    risk = data.get("risk_label", "-")
    started = data.get("started_at", "")
    finished = data.get("finished_at", "")
    elapsed = data.get("elapsed_seconds", 0)
    matches = data.get("matches", []) or []
    ai = data.get("ai_detection") or {}
    cite = data.get("citation_stats") or {}

    lines: list[str] = []
    sep = "=" * 64
    sub = "-" * 64
    lines.append(sep)
    lines.append("  DoubleCheck — Laporan Pengecekan Plagiarisme")
    lines.append(sep)
    lines.append("")
    lines.append(f"  Dokumen      : {name}")
    lines.append(f"  Mulai        : {started}")
    lines.append(f"  Selesai      : {finished}")
    lines.append(f"  Durasi       : {elapsed:.1f} detik")
    lines.append(f"  Paragraf     : {flagged} di-flag dari {total}")
    lines.append("")
    lines.append(f"  SKOR KEMIRIPAN : {score:.1f}%   ({risk})")
    lines.append(sub)
    if cite:
        lines.append(
            f"  Sitasi di-strip : inline={cite.get('inline_citations_stripped', 0)}, "
            f"footnote={cite.get('footnote_refs_stripped', 0)}"
        )
        lines.append(
            f"  Kutipan        : direct={cite.get('direct_quotes_stripped', 0)}, "
            f"block={cite.get('block_quotes_stripped', 0)}"
        )
        if cite.get("reduction_pct"):
            lines.append(f"  Reduksi total  : {int(cite['reduction_pct'])}%")
        lines.append(sub)
    if ai:
        lines.append("")
        lines.append("  AI TEXT DETECTION")
        lines.append(sub)
        lines.append(f"  Kemungkinan AI : {int(ai.get('ai_probability', 0) * 100)}%")
        lines.append(f"  Verdict        : {ai.get('verdict', '-')}")
        sigs = ai.get("signals") or {}
        for k, label in [("roberta", "RoBERTa"), ("perplexity_score", "Perplexity"),
                         ("burstiness", "Burstiness"), ("stylometry_score", "Stylometry")]:
            v = sigs.get(k)
            if v is not None:
                lines.append(f"    - {label:<12}: {int(v * 100)}%")
        lines.append(sub)
    lines.append("")
    if not matches:
        lines.append("  Tidak ada match yang terdeteksi.")
        lines.append("  Dokumen ini bersih dari plagiarisme.")
    else:
        lines.append(f"  SUMBER MIRIP ({len(matches)} match)")
        lines.append(sub)
        for i, m in enumerate(matches, 1):
            pct = m.get("score", 0) * 100
            lines.append(f"\n  [{i}] {m.get('source_title', '(tanpa judul)')}  —  {pct:.1f}%")
            lines.append(f"      Tipe   : {m.get('source_type', '-')}")
            url = m.get("source_url")
            if url:
                lines.append(f"      Link   : {url}")
            qt = (m.get("query_text") or "").strip().replace("\n", " ")
            if qt:
                lines.append(f"      Kutipan: {qt[:300]}{'…' if len(qt) > 300 else ''}")
    lines.append("")
    lines.append(sep)
    lines.append(f"  DoubleCheck · {datetime.utcnow().isoformat(timespec='seconds', sep=' ')}")
    lines.append(sep)
    return "\n".join(lines)


def _download_stem(data: dict, fallback: str) -> str:
    """Build a clean download filename stem from a report.
    Strips the original document's file extension so that
    ``report.txt`` doesn't become ``report.txt.pdf`` on download."""
    name = (data.get("display_name") or "").strip() or fallback
    for ext in (".txt", ".docx", ".doc", ".pdf", ".md", ".rtf"):
        if name.lower().endswith(ext):
            name = name[: -len(ext)]
            break
    safe = name.replace("/", "_").replace("\\", "_").replace('"', "'")
    safe = safe.replace("\n", " ").replace("\r", " ")
    return safe or "laporan"


@app.get("/r/{report_id}/txt")
async def report_txt(
    report_id: str,
    sid: str = Depends(get_session_id),
) -> Response:
    data, _ = load_report(report_id, session_id=sid)
    data = enrich_report(data)
    text = _report_text(data)
    stem = _download_stem(data, report_id)
    return Response(
        text.encode("utf-8"),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{stem}.txt"'},
    )


# ---------- PDF export ----------

def _report_pdf(data: dict) -> bytes:
    from fpdf import FPDF

    name = data.get("display_name") or data.get("document_name", "dokumen")
    score = data.get("overall_score", 0) * 100
    flagged = data.get("flagged_paragraphs", 0)
    total = data.get("total_paragraphs", 0)
    risk = data.get("risk_label", "-")
    started = data.get("started_at", "")
    finished = data.get("finished_at", "")
    elapsed = data.get("elapsed_seconds", 0)
    matches = data.get("matches", []) or []
    ai = data.get("ai_detection") or {}
    cite = data.get("citation_stats") or {}

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(40, 40, 60)
    pdf.cell(0, 10, "DoubleCheck", ln=1)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(100, 100, 110)
    pdf.cell(0, 6, "Laporan Pengecekan Plagiarisme", ln=1)
    pdf.ln(2)

    # Header bar
    pdf.set_draw_color(220, 220, 230)
    pdf.set_line_width(0.3)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)

    # Document info
    pdf.set_text_color(30, 30, 40)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, _safe_pdf(name), ln=1)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(110, 110, 120)
    pdf.cell(0, 5, f"Mulai: {started}", ln=1)
    pdf.cell(0, 5, f"Selesai: {finished}    Durasi: {elapsed:.1f}s", ln=1)
    pdf.cell(0, 5, f"Paragraf: {flagged} di-flag dari {total}", ln=1)
    pdf.ln(3)

    # Big score box
    pdf.set_fill_color(245, 246, 250)
    pdf.set_draw_color(220, 220, 230)
    if score > 30:
        rgb = (220, 60, 70)
    elif score > 10:
        rgb = (220, 150, 50)
    else:
        rgb = (40, 170, 100)
    pdf.set_font("Helvetica", "B", 26)
    pdf.set_text_color(*rgb)
    pdf.cell(95, 18, f"  {score:.1f}%", border=1, ln=0, fill=True)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(95, 18, f"  {risk}", border=1, ln=1, fill=True)
    pdf.set_text_color(80, 80, 90)
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(0, 5, "Skor kemiripan keseluruhan", ln=1)
    pdf.ln(3)

    # Citation strip
    if cite and any(cite.get(k) for k in (
        "inline_citations_stripped", "footnote_refs_stripped",
        "direct_quotes_stripped", "block_quotes_stripped"
    )):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 30, 40)
        pdf.cell(0, 6, "Pembersihan Sitasi", ln=1)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 90)
        pdf.cell(0, 5, (
            f"Sitasi di-strip: inline={cite.get('inline_citations_stripped', 0)}, "
            f"footnote={cite.get('footnote_refs_stripped', 0)}"
        ), ln=1)
        pdf.cell(0, 5, (
            f"Kutipan: direct={cite.get('direct_quotes_stripped', 0)}, "
            f"block={cite.get('block_quotes_stripped', 0)}"
        ), ln=1)
        if cite.get("reduction_pct"):
            pdf.cell(0, 5, f"Reduksi total: {int(cite['reduction_pct'])}%", ln=1)
        pdf.ln(2)

    # AI detection
    if ai:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(30, 30, 40)
        pdf.cell(0, 6, "AI Text Detection", ln=1)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(80, 80, 90)
        pdf.cell(0, 5, f"Kemungkinan AI: {int(ai.get('ai_probability', 0) * 100)}%", ln=1)
        pdf.cell(0, 5, f"Verdict: {ai.get('verdict', '-')}", ln=1)
        sigs = ai.get("signals") or {}
        for k, label in [("roberta", "RoBERTa"), ("perplexity_score", "Perplexity"),
                         ("burstiness", "Burstiness"), ("stylometry_score", "Stylometry")]:
            v = sigs.get(k)
            if v is not None:
                pdf.cell(0, 5, f"  - {label}: {int(v * 100)}%", ln=1)
        pdf.ln(2)

    # Matches
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(30, 30, 40)
    pdf.cell(0, 6, f"Sumber Mirip ({len(matches)} match)", ln=1)

    if not matches:
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(80, 80, 90)
        pdf.cell(0, 6, "Tidak ada match yang terdeteksi. Dokumen ini bersih.", ln=1)
    else:
        pdf.set_font("Helvetica", "", 9)
        for i, m in enumerate(matches, 1):
            pct = m.get("score", 0) * 100
            if pct > 70:
                rgb = (220, 60, 70)
            elif pct > 40:
                rgb = (220, 150, 50)
            else:
                rgb = (40, 170, 100)
            pdf.set_text_color(*rgb)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(0, 6, f"[{i}] {pct:.1f}%  —  {_safe_pdf(m.get('source_title', ''))[:70]}", ln=1)
            pdf.set_text_color(100, 100, 110)
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(0, 4, f"   Tipe: {m.get('source_type', '-')}", ln=1)
            url = m.get("source_url")
            if url:
                pdf.cell(0, 4, f"   Link: {_safe_pdf(url)[:80]}", ln=1)
            qt = (m.get("query_text") or "").strip().replace("\n", " ")
            if qt:
                pdf.set_text_color(80, 80, 90)
                pdf.set_font("Helvetica", "I", 8)
                pdf.multi_cell(0, 4, f"   {_safe_pdf(qt[:280])}{'...' if len(qt) > 280 else ''}")
            pdf.ln(1)

    # Footer
    pdf.ln(4)
    pdf.set_draw_color(220, 220, 230)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.set_text_color(150, 150, 160)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 5, f"DoubleCheck · {datetime.utcnow().isoformat(timespec='seconds', sep=' ')}", ln=1)

    out = pdf.output(dest="S")
    return bytes(out)


def _safe_pdf(s: str) -> str:
    """Strip non-latin-1 chars for Helvetica core font."""
    if not s:
        return ""
    return s.encode("latin-1", "replace").decode("latin-1")


@app.get("/r/{report_id}/pdf")
async def report_pdf(
    report_id: str,
    sid: str = Depends(get_session_id),
) -> Response:
    data, _ = load_report(report_id, session_id=sid)
    data = enrich_report(data)
    pdf_bytes = _report_pdf(data)
    stem = _download_stem(data, report_id)
    return Response(
        pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{stem}.pdf"'},
    )


# ---------- DOCX export ----------

def _report_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = data.get("display_name") or data.get("document_name", "dokumen")
    score = data.get("overall_score", 0) * 100
    flagged = data.get("flagged_paragraphs", 0)
    total = data.get("total_paragraphs", 0)
    risk = data.get("risk_label", "-")
    started = data.get("started_at", "")
    finished = data.get("finished_at", "")
    elapsed = data.get("elapsed_seconds", 0)
    matches = data.get("matches", []) or []
    ai = data.get("ai_detection") or {}
    cite = data.get("citation_stats") or {}

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Title
    title = doc.add_heading("DoubleCheck", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("Laporan Pengecekan Plagiarisme")
    for run in sub.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x80)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Document info
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(13)
    for label, value in [
        ("Mulai", started),
        ("Selesai", finished),
        ("Durasi", f"{elapsed:.1f} detik"),
        ("Paragraf", f"{flagged} di-flag dari {total}"),
    ]:
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        para.add_run(value)
        for run in para.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Score box
    score_p = doc.add_paragraph()
    if score > 30:
        rgb = RGBColor(0xC0, 0x30, 0x40)
    elif score > 10:
        rgb = RGBColor(0xC0, 0x80, 0x20)
    else:
        rgb = RGBColor(0x20, 0x90, 0x60)
    score_run = score_p.add_run(f"SKOR KEMIRIPAN: {score:.1f}%    ")
    score_run.bold = True
    score_run.font.size = Pt(18)
    score_run.font.color.rgb = rgb
    risk_run = score_p.add_run(f"({risk})")
    risk_run.font.size = Pt(12)
    risk_run.font.color.rgb = rgb

    # Citation strip
    if cite and any(cite.get(k) for k in (
        "inline_citations_stripped", "footnote_refs_stripped",
        "direct_quotes_stripped", "block_quotes_stripped"
    )):
        doc.add_heading("Pembersihan Sitasi", level=2)
        for k, v in [
            ("Sitasi inline di-strip", cite.get("inline_citations_stripped", 0)),
            ("Footnote di-strip", cite.get("footnote_refs_stripped", 0)),
            ("Kutipan langsung di-strip", cite.get("direct_quotes_stripped", 0)),
            ("Kutipan block di-strip", cite.get("block_quotes_stripped", 0)),
        ]:
            doc.add_paragraph(f"{k}: {v}")
        if cite.get("reduction_pct"):
            doc.add_paragraph(f"Reduksi total: {int(cite['reduction_pct'])}%")

    # AI detection
    if ai:
        doc.add_heading("AI Text Detection", level=2)
        doc.add_paragraph(f"Kemungkinan AI: {int(ai.get('ai_probability', 0) * 100)}%")
        doc.add_paragraph(f"Verdict: {ai.get('verdict', '-')}")
        sigs = ai.get("signals") or {}
        for k, label in [("roberta", "RoBERTa"), ("perplexity_score", "Perplexity"),
                         ("burstiness", "Burstiness"), ("stylometry_score", "Stylometry")]:
            v = sigs.get(k)
            if v is not None:
                doc.add_paragraph(f"  • {label}: {int(v * 100)}%")

    # Matches
    doc.add_heading(f"Sumber Mirip ({len(matches)} match)", level=2)
    if not matches:
        doc.add_paragraph("Tidak ada match yang terdeteksi. Dokumen ini bersih.")
    else:
        for i, m in enumerate(matches, 1):
            pct = m.get("score", 0) * 100
            head = doc.add_paragraph()
            r = head.add_run(f"[{i}] {m.get('source_title', '(tanpa judul)')}  —  {pct:.1f}%")
            r.bold = True
            if pct > 70:
                r.font.color.rgb = RGBColor(0xC0, 0x30, 0x40)
            elif pct > 40:
                r.font.color.rgb = RGBColor(0xC0, 0x80, 0x20)
            else:
                r.font.color.rgb = RGBColor(0x20, 0x90, 0x60)
            doc.add_paragraph(f"Tipe: {m.get('source_type', '-')}")
            url = m.get("source_url")
            if url:
                doc.add_paragraph(f"Link: {url}")
            qt = (m.get("query_text") or "").strip()
            if qt:
                p = doc.add_paragraph()
                run = p.add_run(qt[:600] + ("..." if len(qt) > 600 else ""))
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x70)
            doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"DoubleCheck · {datetime.utcnow().isoformat(timespec='seconds', sep=' ')}")
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x90, 0x90, 0xA0)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.get("/r/{report_id}/docx")
async def report_docx(
    report_id: str,
    sid: str = Depends(get_session_id),
) -> Response:
    data, _ = load_report(report_id, session_id=sid)
    data = enrich_report(data)
    docx_bytes = _report_docx(data)
    stem = _download_stem(data, report_id)
    return Response(
        docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{stem}.docx"'},
    )


@app.get("/api/health")
async def health() -> dict:
    eng = get_engine()
    return {
        "status": "ok",
        "version": APP_VERSION,
        "corpus_size": len(eng.corpus),
        "semantic_loaded": _semantic_loaded,
        "ensemble_loaded": _ensemble_loaded,
        "cross_encoder_loaded": _cross_encoder_loaded,
        "ai_detector_loaded": _ai_detector_loaded,
        "primary_encoder": (
            "ONNX Hybrid INT8" if _semantic_loaded and not _ensemble_loaded
            else "ONNX Hybrid INT8 + PyTorch mpnet" if _ensemble_loaded
            else "PyTorch (not yet loaded)"
        ),
        "onnx_model": str(ONNX_HYBRID_PATH.relative_to(ROOT)) if ONNX_HYBRID_PATH.exists() else "missing",
    }


@app.post("/api/check")
async def check(
    request: Request,
    file: UploadFile = File(...),
    semantic: bool = Form(False),
    ensemble: bool = Form(False),
    cross_encoder: bool = Form(False),
    strip_citations: bool = Form(True),
    detect_ai: bool = Form(False),
    sid: str = Depends(get_session_id),
) -> dict:
    if not file.filename:
        raise HTTPException(400, "no file")

    # Validate file
    safe_name = Path(file.filename).name
    ext = safe_name.split(".")[-1].lower() if "." in safe_name else ""
    if ext not in ("pdf", "docx", "txt"):
        raise HTTPException(400, f"unsupported file type: .{ext}")

    # Read & check size
    contents = await file.read()
    if len(contents) > MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(400, f"file too large (max {MAX_UPLOAD_MB}MB)")

    # Scope all paths to the session prefix so other users can't see this upload
    sid_prefix = session_prefix(sid)
    stamp = int(time.time() * 1000)
    # Saved file: <ts>_<sid8>_<safename>  (also acts as the public filename)
    saved = UPLOAD_DIR / f"{stamp}_{sid_prefix}_{safe_name}"
    saved.write_bytes(contents)
    # Report ID: <ts>_<sid8>_<safename.stem>
    report_id = f"{stamp}_{sid_prefix}_{Path(safe_name).stem}"

    try:
        engine = get_engine()
        if ensemble:
            ensure_ensemble_loaded()
        elif semantic:
            ensure_semantic_loaded()
        if cross_encoder:
            if not ensemble:
                ensure_semantic_loaded()
            ensure_cross_encoder_loaded()
        if detect_ai:
            ensure_ai_detector_loaded()

        result = engine.check(
            saved,
            use_semantic=semantic or ensemble,
            use_cross_encoder=cross_encoder,
            strip_citations=strip_citations,
            use_ensemble=ensemble,
        )

        result_dict = result.to_dict()
        # Override document_name with our session-scoped filename so the
        # display in /riwayat can strip both the timestamp and session prefix
        result_dict["document_name"] = f"{stamp}_{sid_prefix}_{safe_name}"

        if detect_ai:
            try:
                from core.ai_detector_v2 import detect_ai_enhanced
                from core.parser import parse_any
                text = parse_any(saved)
                ai = detect_ai_enhanced(text, use_perplexity=True, use_stylometry=True)
                result_dict["ai_detection"] = {
                    "ai_probability": ai.ai_probability,
                    "human_probability": ai.human_probability,
                    "verdict": ai.verdict,
                    "confidence": ai.confidence,
                    "model_name": ai.model_name,
                    "signals": ai.signals,
                    "per_paragraph": ai.per_paragraph,
                }
            except Exception as e:  # noqa: BLE001
                result_dict["ai_detection"] = {"error": str(e)}

        # Save JSON report
        report_path = REPORT_DIR / f"{report_id}.json"
        report_path.write_text(
            json.dumps(result_dict, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        return {"id": report_id, **result_dict}
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(500, f"check failed: {e}")


@app.post("/api/report", response_class=HTMLResponse)
async def report_legacy(
    file: UploadFile = File(...),
    semantic: bool = Form(False),
    ensemble: bool = Form(False),
    cross_encoder: bool = Form(False),
    strip_citations: bool = Form(True),
    detect_ai: bool = Form(False),
) -> HTMLResponse:
    """Legacy endpoint: returns plain HTML report (for backward compat)."""
    payload = await check(file, semantic, ensemble, cross_encoder, strip_citations, detect_ai)
    eng = get_engine()
    from core.report import CheckResult, CitationStats, Match
    result = CheckResult(
        document_name=payload.get("document_name", ""),
        total_paragraphs=payload.get("total_paragraphs", 0),
        matches=[Match(**{k: m.get(k) for k in
            ["query_text","matched_text","score","source_title","source_url","source_id","source_type","preview_html"]})
            for m in payload.get("matches", [])],
        overall_score=payload.get("overall_score", 0),
        flagged_paragraphs=payload.get("flagged_paragraphs", 0),
        started_at=payload.get("started_at", ""),
        finished_at=payload.get("finished_at", ""),
        elapsed_seconds=payload.get("elapsed_seconds", 0),
        corpus_size=payload.get("corpus_size", 0),
        citation_stats=CitationStats(**{k: v for k, v in (payload.get("citation_stats") or {}).items()
            if k in CitationStats.__dataclass_fields__}),
    )
    return HTMLResponse(eng.report_html(result))


# -----------------------------------------------------------------------
# Maintenance: periodic cleanup of old session reports
# -----------------------------------------------------------------------
def cleanup_old_reports(max_age_days: int = 7) -> dict:
    """Delete reports and uploaded files older than `max_age_days`.

    Filesystem layout: <timestamp>_<session_prefix>_<name>.[json|txt|pdf|docx].
    The timestamp is the unix-millis at upload time, so we can simply
    parse the first underscore-delimited part and compare to now.
    Returns a stats dict for the caller to log.
    """
    import time as _time
    cutoff_ms = int((_time.time() - max_age_days * 86400) * 1000)
    deleted_reports = 0
    deleted_uploads = 0
    freed_bytes = 0

    for path in REPORT_DIR.glob("*.json"):
        stem = path.stem
        first = stem.split("_", 1)[0] if "_" in stem else ""
        if first.isdigit() and int(first) < cutoff_ms:
            size = path.stat().st_size
            path.unlink()
            deleted_reports += 1
            freed_bytes += size

    for path in UPLOAD_DIR.glob("*"):
        if path.is_file():
            stem = path.stem
            first = stem.split("_", 1)[0] if "_" in stem else ""
            if first.isdigit() and int(first) < cutoff_ms:
                size = path.stat().st_size
                path.unlink()
                deleted_uploads += 1
                freed_bytes += size

    return {
        "deleted_reports": deleted_reports,
        "deleted_uploads": deleted_uploads,
        "freed_mb": round(freed_bytes / 1024 / 1024, 2),
        "max_age_days": max_age_days,
    }


@app.post("/api/admin/cleanup")
async def admin_cleanup(
    max_age_days: int = 7,
) -> dict:
    """Trigger cleanup of old session-scoped reports.

    Idempotent and safe — only touches files older than the cutoff.
    Exposed without auth (admin endpoint) for now; restrict via firewall
    or add token check before production.
    """
    return cleanup_old_reports(max_age_days=max_age_days)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8200)
